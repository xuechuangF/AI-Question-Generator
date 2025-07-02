import os
import json
import asyncio
import aiohttp
from typing import List, Dict, Any, Optional, Tuple
import PyPDF2
import docx
from dataclasses import dataclass, asdict
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
from tqdm import tqdm
import nest_asyncio
import time

# 在Colab中启用嵌套的事件循环
nest_asyncio.apply()

@dataclass
class KnowledgePoint:
    """知识点数据结构"""
    id: int
    title: str
    summary: str
    context_ref: str = ""
    key_formulas: List[str] = None
    key_terms: List[str] = None
    difficulty_level: str = "基础"  # 新增：基础/进阶/高级
    knowledge_type: str = "概念定义"  # 新增：知识点类型

    def __post_init__(self):
        if self.key_formulas is None:
            self.key_formulas = []
        if self.key_terms is None:
            self.key_terms = []

@dataclass
class Question:
    """题目数据结构"""
    id: int
    knowledge_point_id: int
    question: str
    options: Dict[str, str]
    correct_answer: str
    explanation: str
    difficulty: str = "medium"
    question_type: str = "基础理解"  
    related_knowledge_points: List[int] = None  

    def __post_init__(self):
        if self.related_knowledge_points is None:
            self.related_knowledge_points = [self.knowledge_point_id]

class DocumentParser:
    """文档解析器，支持PDF和Word"""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """从PDF中提取文本"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PDF解析错误: {e}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """从Word文档中提取文本"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
        except Exception as e:
            print(f"Word文档解析错误: {e}")
        return text

    @staticmethod
    def parse_document(file_path: str) -> str:
        """根据文件类型解析文档"""
        file_path_lower = file_path.lower()
        
        if file_path_lower.endswith('.pdf'):
            return DocumentParser.extract_text_from_pdf(file_path)
        elif file_path_lower.endswith(('.docx', '.doc')):
            return DocumentParser.extract_text_from_docx(file_path)
        elif file_path_lower.endswith(('.txt', '.md', '.markdown')):
            # 添加对txt和markdown文件的支持
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except UnicodeDecodeError:
                # 如果UTF-8解码失败，尝试其他编码
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
        else:
            raise ValueError("不支持的文件格式。请使用PDF、Word、TXT或Markdown文档。")

class TextChunker:
    """文本分块器"""

    def __init__(self, quality_level: str = "中等"):
        # 根据质量档位设置参数 - 更加精细的区分
        quality_configs = {
            "简约": {"max_chunk_size": 8000, "overlap_size": 200, "min_chunk_size": 3000},
            "中等": {"max_chunk_size": 4000, "overlap_size": 500, "min_chunk_size": 1500},
            "较细致": {"max_chunk_size": 2500, "overlap_size": 600, "min_chunk_size": 1000},
            "细致": {"max_chunk_size": 1800, "overlap_size": 700, "min_chunk_size": 800},
            "精细": {"max_chunk_size": 1200, "overlap_size": 800, "min_chunk_size": 600},
        }

        config = quality_configs.get(quality_level, quality_configs["中等"])
        self.max_chunk_size = config["max_chunk_size"]
        self.overlap_size = config["overlap_size"]
        self.min_chunk_size = config["min_chunk_size"]  # 新增：最小分块大小
        self.quality_level = quality_level

    def chunk_text(self, text: str) -> List[Tuple[str, Dict[str, Any]]]:
        """
        将文本分块，保持语义完整性
        返回: [(chunk_text, metadata), ...]
        """
        # 根据质量档位选择不同的分块策略
        if self.quality_level in ["简约", "中等"]:
            return self._chunk_by_paragraphs(text)
        else:
            return self._chunk_by_sentences(text)  

    def _chunk_by_paragraphs(self, text: str) -> List[Tuple[str, Dict[str, Any]]]:
        """按段落分块（适用于简约和中等档位）"""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        chunk_index = 0

        for i, para in enumerate(paragraphs):
            if len(current_chunk) + len(para) < self.max_chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk and len(current_chunk) >= self.min_chunk_size:
                    chunks.append((
                        current_chunk.strip(),
                        {
                            "chunk_index": chunk_index,
                            "chunk_type": "paragraph_based",
                            "start_paragraph": i - len(current_chunk.split('\n\n')) + 1,
                            "end_paragraph": i - 1
                        }
                    ))
                    chunk_index += 1

                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + para + "\n\n"

        if current_chunk and len(current_chunk) >= self.min_chunk_size:
            chunks.append((
                current_chunk.strip(),
                {
                    "chunk_index": chunk_index,
                    "chunk_type": "paragraph_based",
                    "start_paragraph": len(paragraphs) - len(current_chunk.split('\n\n')),
                    "end_paragraph": len(paragraphs) - 1
                }
            ))

        return chunks

    def _chunk_by_sentences(self, text: str) -> List[Tuple[str, Dict[str, Any]]]:
        """按句子分块（适用于较细致、细致、精细档位）"""
        import re
        # 更精细的句子分割
        sentences = re.split(r'[。！？；\n]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        current_chunk = ""
        chunk_index = 0

        for i, sentence in enumerate(sentences):
            if len(current_chunk) + len(sentence) < self.max_chunk_size:
                current_chunk += sentence + "。"
            else:
                if current_chunk and len(current_chunk) >= self.min_chunk_size:
                    chunks.append((
                        current_chunk.strip(),
                        {
                            "chunk_index": chunk_index,
                            "chunk_type": "sentence_based",
                            "start_sentence": i - current_chunk.count('。') + 1,
                            "end_sentence": i - 1,
                            "granularity": "fine"
                        }
                    ))
                    chunk_index += 1

                # 保留更多重叠内容以保持语义连贯性
                overlap_sentences = current_chunk.split('。')[-3:]  # 保留最后3句
                current_chunk = '。'.join(overlap_sentences) + "。" + sentence + "。"

        if current_chunk and len(current_chunk) >= self.min_chunk_size:
            chunks.append((
                current_chunk.strip(),
                {
                    "chunk_index": chunk_index,
                    "chunk_type": "sentence_based",
                    "start_sentence": len(sentences) - current_chunk.count('。'),
                    "end_sentence": len(sentences) - 1,
                    "granularity": "fine"
                }
            ))

        return chunks

    def _get_overlap_text(self, text: str) -> str:
        """获取重叠文本"""
        if len(text) <= self.overlap_size:
            return text

        # 从后往前找，保持完整的句子
        overlap_start = max(0, len(text) - self.overlap_size)
        # 找到句子开始位置
        while overlap_start > 0 and text[overlap_start] not in '.。!！?？':
            overlap_start -= 1

        return text[overlap_start:].lstrip()

class LLMClient:
    """大模型API客户端"""

    def __init__(self, api_key: str, base_url: str = "https://api.deepseek.com"):
        self.api_key = api_key
        self.base_url = base_url
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

    async def call_api(self, prompt: str, max_tokens: int = 2000) -> str:
      """异步调用API"""
      async with aiohttp.ClientSession() as session:
          data = {
              "model": "deepseek-chat",
              "messages": [{"role": "user", "content": prompt}],
              "max_tokens": max_tokens,
              "temperature": 0.7
          }

          try:

              async with session.post(
                  f"{self.base_url}/v1/chat/completions",
                  headers=self.headers,
                  json=data
              ) as response:
                  if response.status != 200:
                      error_text = await response.text()
                      print(f"API错误 (状态码 {response.status}): {error_text}")
                      return ""

                  result = await response.json()

                  # 检查是否有错误
                  if 'error' in result:
                      print(f"API返回错误: {result['error']}")
                      return ""

                  # 检查是否有正确的响应格式
                  if 'choices' not in result or len(result['choices']) == 0:
                      print(f"API响应格式错误: {result}")
                      return ""

                  return result['choices'][0]['message']['content']
          except Exception as e:
              print(f"API调用错误: {e}")
              import traceback
              traceback.print_exc()
              return ""

class KnowledgeExtractor:
    """知识点提取器"""

    def __init__(self, llm_client: LLMClient, quality_level: str = "中等"):
        self.llm_client = llm_client
        self.quality_level = quality_level

        # 根据质量档位设置不同的提取策略
        self.extraction_strategies = {
            "简约": {
                "focus": "只识别最核心、最重要的主要概念和原理",
                "detail_level": "忽略细节和例子，专注于核心要点",
                "knowledge_types": ["概念定义", "原理方法"],
                "max_kp_per_chunk": 3
            },
            "中等": {
                "focus": "识别重要知识点，包含关键概念、原理和方法",
                "detail_level": "包含必要的细节和重要例子",
                "knowledge_types": ["概念定义", "原理方法", "公式计算", "实例应用"],
                "max_kp_per_chunk": 5
            },
            "较细致": {
                "focus": "详细识别知识点，包含重要细节、例子和注意事项",
                "detail_level": "深入挖掘隐含信息和逻辑关系",
                "knowledge_types": ["概念定义", "原理方法", "公式计算", "实例应用", "注意事项"],
                "max_kp_per_chunk": 7
            },
            "细致": {
                "focus": "全面识别知识点，包含所有重要信息、细节、例子、注意事项",
                "detail_level": "识别前提条件、限制条件、应用场景和对比关系",
                "knowledge_types": ["概念定义", "原理方法", "公式计算", "实例应用", "注意事项", "条件限制", "对比分析"],
                "max_kp_per_chunk": 10
            },
            "精细": {
                "focus": "极其详细地识别所有可能的知识点，包含每个概念、细节、例子、注意事项、前提条件等",
                "detail_level": "提取每个重要的细节、补充说明、隐含假设和边界条件",
                "knowledge_types": ["概念定义", "原理方法", "公式计算", "实例应用", "注意事项", "条件限制", "对比分析", "背景信息", "扩展知识"],
                "max_kp_per_chunk": 15
            }
        }

        self.strategy = self.extraction_strategies.get(quality_level, self.extraction_strategies["中等"])

        self.extraction_prompt_template = """你是一个高级学习辅助AI，任务是分析用户提供的学习资料，将其分解成核心知识点并结构化输出。

**质量档位：** {quality_level}
**提取策略：** {extraction_focus}
**细节要求：** {detail_level}
**预期知识点类型：** {knowledge_types}

**注意：** 你正在处理文档的第 {chunk_index} 部分。

**输入：**
{chunk_text}

**你的任务：**
1. **深入理解内容：** 仔细阅读并理解输入文档的核心主题、逻辑结构和关键信息。

2. **识别知识点：** 根据当前质量档位要求，将文档内容分解成独立、自包含的核心知识点单元：
   {specific_instructions}

3. **为每个知识点创建结构化条目：**
   - `id`: 唯一数字序号 (从1开始)
   - `title`: 简洁、准确的标题 (不超过15字)
   - `summary`: 清晰、精炼的摘要，**务必忠实于原文** (3-5句话)
   - `context_ref`: 指向原文关键位置的引用
   - `key_formulas`: 核心公式列表 (用LaTeX表示)
   - `key_terms`: 最核心的1-3个术语
   - `difficulty_level`: 难度等级 ("基础"/"进阶"/"高级")
   - `knowledge_type`: 知识点类型 (从以下选择: {knowledge_types})

**输出要求：**
严格按以下JSON格式输出：
```json
{{
  "knowledge_points": [
    {{
      "id": 1,
      "title": "知识点标题",
      "summary": "知识点的核心摘要...",
      "context_ref": "原文位置描述",
      "key_formulas": ["$formula1$"],
      "key_terms": ["术语A"],
      "difficulty_level": "基础",
      "knowledge_type": "概念定义"
    }}
  ]
}}
```"""

    def _get_specific_instructions(self, quality_level: str) -> str:
        """根据质量档位生成具体的识别指令"""
        instructions = {
            "简约": """
   - 只识别最核心的概念定义和基本原理
   - 每个文本块最多提取3个最重要的知识点
   - 忽略具体例子和细节说明
   - 专注于"是什么"而非"怎么做"
            """,
            "中等": """
   - 识别重要的概念、原理、方法和关键例子
   - 每个文本块提取重要知识点，参考3-5个，以实际为主
   - 包含必要的公式和计算方法
   - 适当包含重要的应用场景
            """,
            "较细致": """
   - 详细识别概念、原理、方法、例子和注意事项
   - 每个文本块提取知识点，参考5-7个，以实际有意义知识点数量为主
   - 深入挖掘隐含的逻辑关系和条件
   - 识别重要的对比和分类信息
   - 包含具体的操作步骤和应用场景
            """,
            "细致": """
   - 全面识别所有重要的知识要素
   - 每个文本块提取知识点，参考7-10个，以实际有意义知识点数量为主
   - 识别前提条件、限制条件和适用范围
   - 提取重要的注意事项和常见误区
   - 包含详细的例子和对比分析
   - 识别知识点之间的关系和依赖
            """,
            "精细": """
   - 极其详细地识别所有可能的知识点
   - 每个文本块提取知识点，参考10-15个，以实际有意义知识点数量为主
   - 提取每个重要的细节和补充说明
   - 识别隐含的假设和边界条件
   - 包含背景信息和扩展知识
   - 详细的步骤分解和多角度分析
   - 识别潜在的应用场景和限制
            """
        }
        return instructions.get(quality_level, instructions["中等"])

    async def extract_from_chunk(self, chunk_text: str, chunk_metadata: Dict) -> List[KnowledgePoint]:
        """从单个文本块中提取知识点"""
        specific_instructions = self._get_specific_instructions(self.quality_level)
        knowledge_types_str = '", "'.join(self.strategy["knowledge_types"])

        prompt = self.extraction_prompt_template.format(
            quality_level=self.quality_level,
            extraction_focus=self.strategy["focus"],
            detail_level=self.strategy["detail_level"],
            knowledge_types=knowledge_types_str,
            chunk_index=chunk_metadata['chunk_index'] + 1,
            chunk_text=chunk_text,
            specific_instructions=specific_instructions
        )

        response = await self.llm_client.call_api(prompt, max_tokens=4000)  # 增加token数

        try:
            json_match = re.search(r'```json\s*(.*?)\s*```', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
            else:
                json_str = response

            data = json.loads(json_str)
            knowledge_points = []

            for kp_data in data.get('knowledge_points', []):
                kp = KnowledgePoint(
                    id=kp_data['id'],
                    title=kp_data['title'],
                    summary=kp_data['summary'],
                    context_ref=kp_data.get('context_ref', ''),
                    key_formulas=kp_data.get('key_formulas', []),
                    key_terms=kp_data.get('key_terms', []),
                    difficulty_level=kp_data.get('difficulty_level', '基础'),
                    knowledge_type=kp_data.get('knowledge_type', '概念定义')
                )
                knowledge_points.append(kp)

            return knowledge_points
        except Exception as e:
            print(f"解析知识点时出错: {e}")
            return []

    async def extract_all(self, chunks: List[Tuple[str, Dict]]) -> List[KnowledgePoint]:
        """从所有文本块中提取知识点"""
        all_knowledge_points = []

        # 并发处理多个块
        tasks = []
        for chunk_text, chunk_metadata in chunks:
            task = self.extract_from_chunk(chunk_text, chunk_metadata)
            tasks.append(task)

        # 使用进度条
        results = []
        for f in tqdm(asyncio.as_completed(tasks), total=len(tasks), desc="提取知识点"):
            result = await f
            results.extend(result)

        return results

class KnowledgePointMerger:
    """知识点合并器"""

    def __init__(self, quality_level: str = "中等"):
        # 根据质量档位设置不同的相似度阈值和去重策略
        configs = {
            "简约": {"similarity_threshold": 0.60, "title_threshold": 0.70, "aggressive_merge": True},
            "中等": {"similarity_threshold": 0.72, "title_threshold": 0.80, "aggressive_merge": False},
            "较细致": {"similarity_threshold": 0.80, "title_threshold": 0.85, "aggressive_merge": False},
            "细致": {"similarity_threshold": 0.85, "title_threshold": 0.90, "aggressive_merge": False},
            "精细": {"similarity_threshold": 0.90, "title_threshold": 0.95, "aggressive_merge": False},
        }

        config = configs.get(quality_level, configs["中等"])
        self.similarity_threshold = config["similarity_threshold"]
        self.title_threshold = config["title_threshold"]
        self.aggressive_merge = config["aggressive_merge"]
        self.quality_level = quality_level

    def merge_knowledge_points(self, knowledge_points: List[KnowledgePoint]) -> List[KnowledgePoint]:
        """合并相似的知识点，使用多层去重策略"""
        if not knowledge_points:
            return []

        print(f"🔄 开始知识点去重合并 (质量档位: {self.quality_level})")
        print(f"   原始知识点数: {len(knowledge_points)}")

        # 第一步：基于标题的精确去重
        deduped_by_title = self._deduplicate_by_title(knowledge_points)
        print(f"   标题去重后: {len(deduped_by_title)}")

        # 第二步：基于内容相似度的合并
        final_kps = self._merge_by_similarity(deduped_by_title)
        print(f"   相似度合并后: {len(final_kps)}")

        # 第三步：质量过滤和优化
        optimized_kps = self._optimize_knowledge_points(final_kps)
        print(f"   质量优化后: {len(optimized_kps)}")

        # 重新编号
        for i, kp in enumerate(optimized_kps):
            kp.id = i + 1

        return optimized_kps

    def _deduplicate_by_title(self, kps: List[KnowledgePoint]) -> List[KnowledgePoint]:
        """基于标题的精确去重"""
        from difflib import SequenceMatcher

        unique_kps = []
        for kp in kps:
            is_duplicate = False
            for existing_kp in unique_kps:
                # 计算标题相似度
                title_sim = SequenceMatcher(None, kp.title, existing_kp.title).ratio()
                if title_sim > self.title_threshold:
                    # 合并到现有知识点
                    existing_kp.summary = self._merge_summaries(existing_kp.summary, kp.summary)
                    existing_kp.key_formulas = list(set(existing_kp.key_formulas + kp.key_formulas))
                    existing_kp.key_terms = list(set(existing_kp.key_terms + kp.key_terms))
                    is_duplicate = True
                    break

            if not is_duplicate:
                unique_kps.append(kp)

        return unique_kps

    def _merge_by_similarity(self, kps: List[KnowledgePoint]) -> List[KnowledgePoint]:
        """基于内容相似度的合并"""
        if len(kps) <= 1:
            return kps

        # 构建文本向量
        texts = [f"{kp.title} {kp.summary} {' '.join(kp.key_terms)}" for kp in kps]

        try:
            vectorizer = TfidfVectorizer(max_features=200, stop_words=None)
            tfidf_matrix = vectorizer.fit_transform(texts)
            similarity_matrix = cosine_similarity(tfidf_matrix)
        except:
            # 如果向量化失败，返回原始列表
            return kps

        # 标记已合并的知识点
        merged = [False] * len(kps)
        merged_kps = []

        for i in range(len(kps)):
            if merged[i]:
                continue

            # 找到相似的知识点
            similar_indices = []
            for j in range(i + 1, len(kps)):
                if not merged[j] and similarity_matrix[i][j] > self.similarity_threshold:
                    # 额外检查：确保知识点类型相同或相关
                    if self._are_types_compatible(kps[i].knowledge_type, kps[j].knowledge_type):
                        similar_indices.append(j)

            # 合并知识点
            if similar_indices and self.aggressive_merge:
                merged_kp = self._merge_multiple([kps[i]] + [kps[j] for j in similar_indices])
                merged_kps.append(merged_kp)
                merged[i] = True
                for j in similar_indices:
                    merged[j] = True
            else:
                merged_kps.append(kps[i])
                merged[i] = True

        return merged_kps

    def _are_types_compatible(self, type1: str, type2: str) -> bool:
        """检查两个知识点类型是否可以合并"""
        # 定义可以合并的知识点类型组合
        compatible_groups = [
            ["概念定义", "背景信息"],
            ["原理方法", "实例应用"],
            ["公式计算", "原理方法"],
            ["注意事项", "条件限制"],
            ["对比分析", "概念定义"]
        ]

        if type1 == type2:
            return True

        for group in compatible_groups:
            if type1 in group and type2 in group:
                return True

        return False

    def _merge_summaries(self, summary1: str, summary2: str) -> str:
        """智能合并两个摘要"""
        sentences1 = [s.strip() for s in summary1.split('。') if s.strip()]
        sentences2 = [s.strip() for s in summary2.split('。') if s.strip()]

        # 去除重复句子
        unique_sentences = []
        all_sentences = sentences1 + sentences2

        for sentence in all_sentences:
            if sentence and not any(self._sentence_similarity(sentence, existing) > 0.8
                                  for existing in unique_sentences):
                unique_sentences.append(sentence)

        # 限制长度，优先保留更有信息量的句子
        if len(unique_sentences) > 5:
            unique_sentences = sorted(unique_sentences, key=len, reverse=True)[:5]

        return '。'.join(unique_sentences) + '。'

    def _sentence_similarity(self, sent1: str, sent2: str) -> float:
        """计算两个句子的相似度"""
        from difflib import SequenceMatcher
        return SequenceMatcher(None, sent1, sent2).ratio()

    def _optimize_knowledge_points(self, kps: List[KnowledgePoint]) -> List[KnowledgePoint]:
        """优化知识点质量"""
        optimized = []

        for kp in kps:
            # 过滤掉质量过低的知识点
            if len(kp.title) < 3 or len(kp.summary) < 10:
                continue

            # 清理和优化内容
            kp.title = kp.title.strip()
            kp.summary = kp.summary.strip()
            kp.key_formulas = [f.strip() for f in kp.key_formulas if f.strip()]
            kp.key_terms = [t.strip() for t in kp.key_terms if t.strip()]

            optimized.append(kp)

        return optimized

    def _merge_multiple(self, kps: List[KnowledgePoint]) -> KnowledgePoint:
        """合并多个知识点"""
        # 选择最长的标题
        title = max([kp.title for kp in kps], key=len)

        # 智能合并摘要
        summaries = [kp.summary for kp in kps]
        merged_summary = self._merge_summaries(summaries[0],
                                             '。'.join(summaries[1:]) if len(summaries) > 1 else "")

        # 合并其他字段
        context_refs = [kp.context_ref for kp in kps if kp.context_ref]
        key_formulas = list(set(sum([kp.key_formulas for kp in kps], [])))
        key_terms = list(set(sum([kp.key_terms for kp in kps], [])))

        # 选择最高的难度级别
        difficulty_levels = [kp.difficulty_level for kp in kps]
        difficulty_order = {"基础": 1, "进阶": 2, "高级": 3}
        final_difficulty = max(difficulty_levels, key=lambda x: difficulty_order.get(x, 1))

        return KnowledgePoint(
            id=kps[0].id,
            title=title,
            summary=merged_summary,
            context_ref='; '.join(context_refs),
            key_formulas=key_formulas,
            key_terms=key_terms,
            difficulty_level=final_difficulty,
            knowledge_type=kps[0].knowledge_type  # 使用第一个的类型
        )

class QuestionGenerator:
    """增强版题目生成器"""

    def __init__(self, llm_client: LLMClient, quality_level: str = "中等"):
        self.llm_client = llm_client
        self.quality_level = quality_level

        # 根据质量档位配置题目生成策略
        self.question_configs = {
            "简约": {"basic_per_kp": 1, "fusion_ratio": 0.1, "advanced_ratio": 0.1},
            "中等": {"basic_per_kp": 1, "fusion_ratio": 0.2, "advanced_ratio": 0.2},
            "较细致": {"basic_per_kp": 1, "fusion_ratio": 0.3, "advanced_ratio": 0.3},
            "细致": {"basic_per_kp": 1, "fusion_ratio": 0.4, "advanced_ratio": 0.4},
            "精细": {"basic_per_kp": 1, "fusion_ratio": 0.5, "advanced_ratio": 0.5}
        }

        self.basic_question_template = """请严格根据以下提供的单一知识点摘要，生成一道高质量四选项单项选择题。

知识点信息:
- 标题: {title}
- 难度: {difficulty_level}
- 类型: {knowledge_type}
- 核心摘要: {summary}
{formulas_text}

题目要求:
1. 题目要准确测试对该知识点的{test_focus}
2. 正确答案必须完全基于摘要内容
3. 三个错误选项要有迷惑性但明显错误
4. 题目难度为: {target_difficulty}
5. 如果涉及公式，使用LaTeX格式 $formula$
6. **重要：正确答案随机分布在A、B、C、D中**

输出格式（严格JSON）:
```json
{{
  "question": "题目内容",
  "options": {{
    "A": "选项A内容",
    "B": "选项B内容",
    "C": "选项C内容",
    "D": "选项D内容"
  }},
  "correct_answer": "A/B/C/D",
  "explanation": "答案解释",
  "difficulty": "{target_difficulty}",
  "question_type": "基础理解"
}}
```"""

        self.fusion_question_template = """请根据以下多个相关知识点，生成一道融合性四选项单项选择题，要求综合运用多个知识点。

相关知识点:
{knowledge_points_info}

题目要求:
1. 题目需要综合运用上述多个知识点才能正确回答
2. 题型为: {question_type}
3. 难度等级: {target_difficulty}
4. 三个错误选项要基于部分知识点但结论错误
5. 正确答案随机分布在A、B、C、D中

输出格式（严格JSON）:
```json
{{
  "question": "题目内容",
  "options": {{
    "A": "选项A内容",
    "B": "选项B内容",
    "C": "选项C内容",
    "D": "选项D内容"
  }},
  "correct_answer": "A/B/C/D",
  "explanation": "综合解释，说明涉及的多个知识点",
  "difficulty": "{target_difficulty}",
  "question_type": "{question_type}",
  "related_knowledge_points": {related_kp_ids}
}}
```"""

    async def generate_basic_question(self, kp: KnowledgePoint, question_id: int, target_difficulty: str = "medium") -> Optional[Question]:
        """为单个知识点生成基础题目"""
        import random

        # 根据知识点类型确定测试重点
        test_focus_map = {
            "概念定义": "概念理解和定义记忆",
            "原理方法": "原理理解和方法应用",
            "公式计算": "公式理解和计算能力",
            "实例应用": "实际应用和案例分析",
            "注意事项": "注意事项和限制条件的理解"
        }

        test_focus = test_focus_map.get(kp.knowledge_type, "核心理解")

        formulas_text = ""
        if kp.key_formulas:
            formulas_text = f"相关公式: {', '.join(kp.key_formulas)}"

        prompt = self.basic_question_template.format(
            title=kp.title,
            difficulty_level=kp.difficulty_level,
            knowledge_type=kp.knowledge_type,
            summary=kp.summary,
            formulas_text=formulas_text,
            test_focus=test_focus,
            target_difficulty=target_difficulty
        )

        response = await self.llm_client.call_api(prompt)
        return self._parse_question_response(response, question_id, kp.id)

    async def generate_fusion_question(self, kps: List[KnowledgePoint], question_id: int, question_type: str) -> Optional[Question]:
        """生成融合多个知识点的题目"""

        # 构建知识点信息
        kp_info_list = []
        for i, kp in enumerate(kps, 1):
            kp_info = f"{i}. {kp.title}: {kp.summary}"
            if kp.key_formulas:
                kp_info += f" (公式: {', '.join(kp.key_formulas)})"
            kp_info_list.append(kp_info)

        knowledge_points_info = "\n".join(kp_info_list)
        related_kp_ids = [kp.id for kp in kps]

        # 根据涉及知识点数量和类型确定难度
        target_difficulty = "hard" if len(kps) >= 3 else "medium"

        prompt = self.fusion_question_template.format(
            knowledge_points_info=knowledge_points_info,
            question_type=question_type,
            target_difficulty=target_difficulty,
            related_kp_ids=related_kp_ids
        )

        response = await self.llm_client.call_api(prompt)
        question = self._parse_question_response(response, question_id, kps[0].id)

        if question:
            question.related_knowledge_points = related_kp_ids

        return question

    def _parse_question_response(self, response: str, question_id: int, kp_id: int) -> Optional[Question]:
        """解析题目生成响应"""
        import random
        try:
            # 提取JSON内容
            json_match = re.search(r'```json\s*(.*?)\s*```', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
            else:
                json_str = response

            json_str = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)

            json_str = json_str.replace('\b', '').replace('\f', '').replace('\v', '')

            json_str = json_str.replace('\\', '\\\\').replace('\\"', '"')

            data = json.loads(json_str)

            # 随机化答案位置
            options = data['options']
            correct_answer = data['correct_answer']

            correct_content = options[correct_answer]

            option_contents = list(options.values())
            random.shuffle(option_contents)
            new_options = {}
            new_correct_answer = None
            option_keys = ['A', 'B', 'C', 'D']

            for i, content in enumerate(option_contents):
                new_options[option_keys[i]] = content
                if content == correct_content:
                    new_correct_answer = option_keys[i]

            question = Question(
                id=question_id,
                knowledge_point_id=kp_id,
                question=data['question'],
                options=new_options,
                correct_answer=new_correct_answer,
                explanation=data['explanation'],
                difficulty=data.get('difficulty', 'medium')
            )

            question.question_type = data.get('question_type', '基础理解')
            question.related_knowledge_points = data.get('related_knowledge_points', [kp_id])

            return question

        except Exception as e:
            print(f"生成题目时出错: {e}")
            return None

    async def generate_all(self, knowledge_points: List[KnowledgePoint]) -> List[Question]:
        """为所有知识点生成完整的题目集"""
        config = self.question_configs.get(self.quality_level, self.question_configs["中等"])
        questions = []
        question_id = 1

        print(f"📊 题目生成策略 - 基础题:{config['basic_per_kp']}题/知识点, 融合题比例:{config['fusion_ratio']:.0%}, 高难度比例:{config['advanced_ratio']:.0%}")

        # 1. 为每个知识点生成基础题目
        basic_tasks = []
        for kp in knowledge_points:
            for _ in range(config['basic_per_kp']):
                difficulty = "hard" if kp.difficulty_level == "高级" else "medium" if kp.difficulty_level == "进阶" else "easy"
                task = self.generate_basic_question(kp, question_id, difficulty)
                basic_tasks.append(task)
                question_id += 1

        # 生成基础题目
        for f in tqdm(asyncio.as_completed(basic_tasks), total=len(basic_tasks), desc="生成基础题目"):
            question = await f
            if question:
                questions.append(question)

        # 2. 生成融合题目
        fusion_count = int(len(knowledge_points) * config['fusion_ratio'])
        fusion_types = [
            "因果推理型", "综合关联型", "最优方案型", "情景应用题",
            "对比分析型", "综合判断型", "系统分析型"
        ]

        fusion_tasks = []
        for _ in range(fusion_count):
            # 随机选择2-4个相关知识点
            import random
            selected_kps = random.sample(knowledge_points, min(random.randint(2, 4), len(knowledge_points)))
            question_type = random.choice(fusion_types)

            task = self.generate_fusion_question(selected_kps, question_id, question_type)
            fusion_tasks.append(task)
            question_id += 1

        # 生成融合题目
        for f in tqdm(asyncio.as_completed(fusion_tasks), total=len(fusion_tasks), desc="生成融合题目"):
            question = await f
            if question:
                questions.append(question)

        # 3. 生成额外的高难度题目
        advanced_count = int(len(knowledge_points) * config['advanced_ratio'])
        advanced_tasks = []

        # 选择难度较高的知识点
        advanced_kps = [kp for kp in knowledge_points if kp.difficulty_level in ["进阶", "高级"]]
        if not advanced_kps:
            advanced_kps = knowledge_points  # 如果没有高难度知识点，使用所有知识点

        for _ in range(advanced_count):
            import random
            kp = random.choice(advanced_kps)
            task = self.generate_basic_question(kp, question_id, "hard")
            advanced_tasks.append(task)
            question_id += 1

        # 生成高难度题目
        for f in tqdm(asyncio.as_completed(advanced_tasks), total=len(advanced_tasks), desc="生成高难度题目"):
            question = await f
            if question:
                questions.append(question)

        # 按ID排序
        questions.sort(key=lambda q: q.id)
        return questions

class NoteToQuizGenerator:
    """笔记生题器主类"""

    def __init__(self, api_key: str, quality_level: str = "中等"):
        self.llm_client = LLMClient(api_key)
        self.chunker = TextChunker(quality_level)
        self.extractor = KnowledgeExtractor(self.llm_client, quality_level)
        self.merger = KnowledgePointMerger(quality_level)
        self.generator = QuestionGenerator(self.llm_client, quality_level)

    async def process_document(self, file_path: str) -> Tuple[List[KnowledgePoint], List[Question]]:
        """处理文档并生成题目"""
        print("📄 正在解析文档...")
        text = DocumentParser.parse_document(file_path)

        print("✂️ 正在分块处理...")
        chunks = self.chunker.chunk_text(text)
        print(f"文档被分成 {len(chunks)} 个块")

        print("🔍 正在提取知识点...")
        raw_knowledge_points = await self.extractor.extract_all(chunks)
        print(f"提取到 {len(raw_knowledge_points)} 个原始知识点")

        print("🔗 正在合并相似知识点...")
        merged_knowledge_points = self.merger.merge_knowledge_points(raw_knowledge_points)
        print(f"合并后剩余 {len(merged_knowledge_points)} 个知识点")

        print("📝 正在生成题目...")
        questions = await self.generator.generate_all(merged_knowledge_points)
        print(f"成功生成 {len(questions)} 道题目")

        return merged_knowledge_points, questions

    def save_results(self, knowledge_points: List[KnowledgePoint],
                    questions: List[Question],
                    output_dir: str = "output"):
        """保存结果"""
        os.makedirs(output_dir, exist_ok=True)

        # 保存知识点
        kp_data = [asdict(kp) for kp in knowledge_points]
        with open(f"{output_dir}/knowledge_points.json", 'w', encoding='utf-8') as f:
            json.dump(kp_data, f, ensure_ascii=False, indent=2)

        # 保存题目
        q_data = [asdict(q) for q in questions]
        with open(f"{output_dir}/questions.json", 'w', encoding='utf-8') as f:
            json.dump(q_data, f, ensure_ascii=False, indent=2)

        # 生成可打印的题目文档
        with open(f"{output_dir}/quiz.txt", 'w', encoding='utf-8') as f:
            f.write("=== 生成的测验题目 ===\n\n")
            for q in questions:
                f.write(f"题目 {q.id}. {q.question}\n")
                for opt, content in q.options.items():
                    f.write(f"  {opt}. {content}\n")
                f.write(f"\n")

            f.write("\n\n=== 答案和解释 ===\n\n")
            for q in questions:
                f.write(f"题目 {q.id}: {q.correct_answer}\n")
                f.write(f"解释: {q.explanation}\n\n")

        print(f"✅ 结果已保存到 {output_dir} 目录")


class InteractiveReviewer:
    """交互式知识点审核器"""

    def review_knowledge_points(self, knowledge_points: List[KnowledgePoint]) -> List[KnowledgePoint]:
        """让用户审核和编辑知识点"""
        print("\n📋 知识点审核")
        print("=" * 50)

        reviewed_kps = []

        for i, kp in enumerate(knowledge_points):
            print(f"\n知识点 {i+1}/{len(knowledge_points)}")
            print(f"标题: {kp.title}")
            print(f"摘要: {kp.summary}")
            if kp.key_formulas:
                print(f"公式: {', '.join(kp.key_formulas)}")

            action = input("\n操作: [K]保留 [E]编辑 [D]删除 ：").upper()

            if action == 'K':
                reviewed_kps.append(kp)
            elif action == 'E':
                new_title = input(f"新标题 (回车保持原标题): ").strip()
                new_summary = input(f"新摘要 (回车保持原摘要): ").strip()

                if new_title:
                    kp.title = new_title
                if new_summary:
                    kp.summary = new_summary

                reviewed_kps.append(kp)
            # D - 删除，不添加到reviewed_kps

        for i, kp in enumerate(reviewed_kps):
            kp.id = i + 1

        print(f"\n✅ 审核完成！保留了 {len(reviewed_kps)} 个知识点")
        return reviewed_kps

# ========== 批量处理功能 ==========

class BatchProcessor:
    """批量处理多个文档"""

    def __init__(self, generator: NoteToQuizGenerator):
        self.generator = generator

    async def process_multiple_documents(self, file_paths: List[str]) -> Dict[str, Tuple[List[KnowledgePoint], List[Question]]]:
        """批量处理多个文档"""
        results = {}

        for file_path in file_paths:
            print(f"\n处理文档: {file_path}")
            try:
                kps, questions = await self.generator.process_document(file_path)
                results[file_path] = (kps, questions)
            except Exception as e:
                print(f"处理 {file_path} 时出错: {e}")
                results[file_path] = ([], [])

        return results

# ========== 题目格式化输出 ==========

class QuizFormatter:
    """题目格式化器"""

    @staticmethod
    def to_html(questions: List[Question], knowledge_points: List[KnowledgePoint]) -> str:
        """生成HTML格式的测验"""
        html = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>知识点测验</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }
        .question {
            margin-bottom: 30px;
            padding: 20px;
            background-color: #f5f5f5;
            border-radius: 8px;
        }
        .question-number {
            font-weight: bold;
            color: #333;
        }
        .options {
            margin-top: 10px;
        }
        .option {
            margin: 5px 0;
            padding: 5px 10px;
        }
        .answer-section {
            margin-top: 50px;
            border-top: 2px solid #ccc;
            padding-top: 20px;
        }
        .answer {
            margin: 10px 0;
            padding: 10px;
            background-color: #e8f5e9;
            border-radius: 5px;
        }
        .formula {
            font-style: italic;
            color: #1976d2;
        }
        .knowledge-point {
            margin: 20px 0;
            padding: 15px;
            background-color: #f0f7ff;
            border-left: 4px solid #2196f3;
        }
    </style>
</head>
<body>
    <h1>知识点测验</h1>

    <h2>知识点概览</h2>
    <div class="knowledge-points">
"""

        # 添加知识点概览
        for kp in knowledge_points:
            html += f"""
        <div class="knowledge-point">
            <h3>{kp.id}. {kp.title}</h3>
            <p>{kp.summary}</p>
        </div>
"""

        html += """
    </div>

    <h2>测验题目</h2>
    <div class="questions">
"""

        
        for q in questions:
            html += f"""
        <div class="question">
            <div class="question-number">题目 {q.id}</div>
            <p>{q.question}</p>
            <div class="options">
"""
            for opt, content in q.options.items():
                html += f'                <div class="option">{opt}. {content}</div>\n'

            html += """            </div>
        </div>
"""

        
        html += """
    </div>

    <div class="answer-section">
        <h2>答案与解析</h2>
"""

        for q in questions:
            html += f"""
        <div class="answer">
            <strong>题目 {q.id}:</strong> {q.correct_answer}<br>
            <strong>解析:</strong> {q.explanation}
        </div>
"""

        html += """
    </div>
</body>
</html>
"""
        return html

    @staticmethod
    def to_markdown(questions: List[Question], knowledge_points: List[KnowledgePoint]) -> str:
        """生成Markdown格式的测验"""
        md = "# 知识点测验\n\n"

        md += "## 知识点概览\n\n"
        for kp in knowledge_points:
            md += f"### {kp.id}. {kp.title}\n\n"
            md += f"{kp.summary}\n\n"
            if kp.key_formulas:
                md += f"**关键公式：** {', '.join(kp.key_formulas)}\n\n"

        md += "## 测验题目\n\n"
        for q in questions:
            md += f"### 题目 {q.id}\n\n"
            md += f"{q.question}\n\n"
            for opt, content in q.options.items():
                md += f"- {opt}. {content}\n"
            md += "\n"

        md += "## 答案与解析\n\n"
        for q in questions:
            md += f"**题目 {q.id}:** {q.correct_answer}\n\n"
            md += f"**解析：** {q.explanation}\n\n"

        return md

class Config:
    """配置管理类"""

    API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
    API_BASE_URL = "https://api.deepseek.com"

    QUALITY_LEVEL = "中等"

    MAX_CONCURRENT_REQUESTS = 5

    OUTPUT_DIR = "output"

    @classmethod
    def from_dict(cls, config_dict: Dict[str, Any]):
        """从字典加载配置"""
        for key, value in config_dict.items():
            if hasattr(cls, key):
                setattr(cls, key, value)

    @classmethod
    def to_dict(cls) -> Dict[str, Any]:
        """导出配置为字典"""
        return {
            "API_KEY": cls.API_KEY,
            "API_BASE_URL": cls.API_BASE_URL,
            "QUALITY_LEVEL": cls.QUALITY_LEVEL,
            "MAX_CONCURRENT_REQUESTS": cls.MAX_CONCURRENT_REQUESTS,
            "OUTPUT_DIR": cls.OUTPUT_DIR
        }

class EnhancedNoteToQuizGenerator(NoteToQuizGenerator):
    """笔记生题器"""

    def __init__(self, config: Config = None):
        if config is None:
            config = Config()

        super().__init__(config.API_KEY, config.QUALITY_LEVEL)  
        self.config = config
        self.chunker = TextChunker(quality_level=config.QUALITY_LEVEL)
        self.merger = KnowledgePointMerger(quality_level=config.QUALITY_LEVEL)
        self.reviewer = InteractiveReviewer()
        self.formatter = QuizFormatter()

    async def process_with_review(self, file_path: str, enable_review: bool = True) -> Tuple[List[KnowledgePoint], List[Question]]:
        """处理文档并可选地进行人工审核"""
        text = DocumentParser.parse_document(file_path)
        chunks = self.chunker.chunk_text(text)
        raw_knowledge_points = await self.extractor.extract_all(chunks)
        merged_knowledge_points = self.merger.merge_knowledge_points(raw_knowledge_points)

        if enable_review:
            merged_knowledge_points = self.reviewer.review_knowledge_points(merged_knowledge_points)

        questions = await self.generator.generate_all(merged_knowledge_points)

        return merged_knowledge_points, questions

    def save_all_formats(self, knowledge_points: List[KnowledgePoint],
                        questions: List[Question],
                        base_name: str = "quiz"):
        """保存为多种格式"""
        output_dir = self.config.OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)

        # JSON格式
        self.save_results(knowledge_points, questions, output_dir)

        # HTML格式
        html_content = self.formatter.to_html(questions, knowledge_points)
        with open(f"{output_dir}/{base_name}.html", 'w', encoding='utf-8') as f:
            f.write(html_content)

        # Markdown格式
        md_content = self.formatter.to_markdown(questions, knowledge_points)
        with open(f"{output_dir}/{base_name}.md", 'w', encoding='utf-8') as f:
            f.write(md_content)

        print(f"✅ 已保存为多种格式到 {output_dir} 目录")


class ColabHelper:
    """Colab环境辅助工具"""

    @staticmethod
    def setup_environment():
        """设置Colab环境"""
        print("🔧 正在设置Colab环境...")

        import subprocess
        subprocess.run(["pip", "install", "-q", "PyPDF2", "python-docx",
                       "scikit-learn", "aiohttp", "nest_asyncio", "tqdm"])

        print("✅ 依赖安装完成")

    @staticmethod
    def upload_file():
        """上传文件到Colab"""
        from google.colab import files
        print("📤 请选择要上传的文档...")
        uploaded = files.upload()

        if uploaded:
            file_name = list(uploaded.keys())[0]
            print(f"✅ 文件 '{file_name}' 上传成功")
            return file_name
        else:
            print("❌ 没有上传文件")
            return None

    @staticmethod
    def download_results(output_dir: str = "output"):
        """下载结果文件"""
        from google.colab import files

        print("📥 正在准备下载文件...")

        import zipfile
        zip_name = "quiz_results.zip"

        with zipfile.ZipFile(zip_name, 'w') as zipf:
            for root, dirs, files_list in os.walk(output_dir):
                for file in files_list:
                    file_path = os.path.join(root, file)
                    zipf.write(file_path)

        files.download(zip_name)
        print("✅ 结果文件已下载")

    @staticmethod
    def interactive_setup():
        """交互式设置"""
        print("🎯 笔记生题器 - 交互式设置")
        print("=" * 50)

        api_key = input("请输入你的DeepSeek API密钥: ").strip()

        print("\n📋 选择题目生成质量档位:")
        print("1. 简约 - 只抓重点知识，生成核心题目")
        print("2. 中等 - 平衡覆盖，适中的题目数量")
        print("3. 较细致 - 覆盖更多知识点，增加融合题")
        print("4. 细致 - 详细覆盖大部分知识点，多种题型")
        print("5. 精细 - 尽可能抓全所有知识点，丰富题型")

        quality_choice = input("请选择质量档位 [1-5，默认2]: ").strip() or "2"

        quality_map = {
            "1": "简约",
            "2": "中等",
            "3": "较细致",
            "4": "细致",
            "5": "精细"
        }

        quality_level = quality_map.get(quality_choice, "中等")
        print(f"已选择质量档位: {quality_level}")

        config = Config()
        config.API_KEY = api_key
        config.QUALITY_LEVEL = quality_level

        return config

async def colab_main():
    """Colab环境的主函数"""
    # 1. 设置环境
    ColabHelper.setup_environment()

    # 2. 交互式配置
    config = ColabHelper.interactive_setup()

    # 3. 上传文件
    file_path = ColabHelper.upload_file()
    if not file_path:
        return

    # 4. 创建生成器
    generator = EnhancedNoteToQuizGenerator(config)

    # 5. 询问是否需要审核
    enable_review = input("\n是否需要人工审核知识点？[y/N]: ").lower() == 'y'

    try:
        # 6. 处理文档
        print("\n🚀 开始处理文档...")
        start_time = time.time()

        knowledge_points, questions = await generator.process_with_review(
            file_path, enable_review
        )

        # 7. 保存结果
        generator.save_all_formats(knowledge_points, questions,
                                 base_name=os.path.splitext(file_path)[0])

        # 8. 显示统计
        elapsed_time = time.time() - start_time
        print(f"\n⏱️ 总用时: {elapsed_time:.2f} 秒")
        print(f"📊 生成了 {len(questions)} 道题目")

        # 9. 下载结果
        if input("\n是否下载结果文件？[Y/n]: ").lower() != 'n':
            ColabHelper.download_results(config.OUTPUT_DIR)

    except Exception as e:
        print(f"\n❌ 处理过程中出错: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    try:
        import google.colab
        IN_COLAB = True
    except ImportError:
        IN_COLAB = False

    if IN_COLAB:
        asyncio.run(colab_main())
    else:
        asyncio.run(main())
